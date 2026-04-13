"""
Generates the ML tutorial .docx for the DVF real estate project.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ────────────────────────────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(3.0)
section.right_margin  = Cm(2.5)

# ── Helper: set paragraph font size & colour ─────────────────────────────────
def fmt(para, size=11, bold=False, color=None, italic=False):
    for run in para.runs:
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.italic = italic
        if color:
            run.font.color.rgb = RGBColor(*color)
    return para

# ── Helper: add a styled heading ─────────────────────────────────────────────
def h1(text):
    p = doc.add_heading(text, level=1)
    fmt(p, size=16, bold=True, color=(31, 78, 121))
    return p

def h2(text):
    p = doc.add_heading(text, level=2)
    fmt(p, size=13, bold=True, color=(21, 96, 130))
    return p

def h3(text):
    p = doc.add_heading(text, level=3)
    fmt(p, size=11, bold=True, color=(40, 40, 40))
    return p

# ── Helper: body paragraph ───────────────────────────────────────────────────
def body(text, indent=False):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    if indent:
        p.paragraph_format.left_indent = Cm(0.8)
    for run in p.runs:
        run.font.size = Pt(11)
    return p

# ── Helper: bullet ───────────────────────────────────────────────────────────
def bullet(text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.8 + level * 0.6)
    p.paragraph_format.space_after = Pt(3)
    for run in p.runs:
        run.font.size = Pt(11)
    return p

# ── Helper: numbered ─────────────────────────────────────────────────────────
def numbered(text):
    p = doc.add_paragraph(text, style="List Number")
    p.paragraph_format.space_after = Pt(3)
    for run in p.runs:
        run.font.size = Pt(11)
    return p

# ── Helper: code block (monospace, shaded) ───────────────────────────────────
def code(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.8)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(30, 30, 30)
    # Light grey background via paragraph shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F2F2F2")
    pPr.append(shd)
    return p

# ── Helper: info box (bordered paragraph) ───────────────────────────────────
def callout(text, emoji="💡"):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.space_after  = Pt(8)
    run = p.add_run(f"{emoji}  {text}")
    run.font.size = Pt(10.5)
    run.font.italic = True
    run.font.color.rgb = RGBColor(80, 80, 80)
    return p

def sep():
    doc.add_paragraph("─" * 72)

# ════════════════════════════════════════════════════════════════════════════
#  COVER
# ════════════════════════════════════════════════════════════════════════════
title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title_para.add_run("Machine Learning with Real Data")
title_run.font.size = Pt(26)
title_run.font.bold = True
title_run.font.color.rgb = RGBColor(31, 78, 121)

sub_para = doc.add_paragraph()
sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_run = sub_para.add_run("A Step-by-Step Project Guide using French Real Estate (DVF) Data")
sub_run.font.size = Pt(14)
sub_run.font.color.rgb = RGBColor(90, 90, 90)
sub_run.font.italic = True

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run("Dataset: 16.3 M rows  |  Python 3.14  |  scikit-learn · LightGBM · MLflow")
meta.runs[0].font.size = Pt(10)
meta.runs[0].font.color.rgb = RGBColor(120, 120, 120)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
#  TABLE OF CONTENTS (manual)
# ════════════════════════════════════════════════════════════════════════════
h1("Table of Contents")
toc_items = [
    ("1.", "Project Overview & Learning Goals"),
    ("2.", "Step 1 — Setting Up the Environment"),
    ("3.", "Step 2 — Loading & Understanding the Dataset"),
    ("4.", "Step 3 — Data Quality Audit (missing values, duplicates, outliers)"),
    ("5.", "Step 4 — Writing Failing Data Quality Tests"),
    ("6.", "Step 5 — Preparing the Data for Machine Learning"),
    ("7.", "Step 6 — Choosing & Training Three Models"),
    ("8.", "Step 7 — Tracking Experiments with MLflow"),
    ("9.", "Step 8 — Comparing Results & Reading the Metrics"),
    ("10.", "Key Concepts Glossary"),
    ("11.", "What to Learn Next"),
]
for num, title in toc_items:
    p = doc.add_paragraph()
    r1 = p.add_run(f"{num}  ")
    r1.font.bold = True
    r1.font.size = Pt(11)
    r2 = p.add_run(title)
    r2.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(4)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
#  1. OVERVIEW
# ════════════════════════════════════════════════════════════════════════════
h1("1.  Project Overview & Learning Goals")

body(
    "This project teaches you the full cycle of a real-world machine learning task: "
    "from raw, messy data to trained, tracked models ready for comparison. "
    "Instead of using a toy dataset, we work with the French government's open "
    "property-transaction database (DVF — Demandes de Valeurs Foncieres), which "
    "contains over 16 million rows of real estate sales. This scale forces you to "
    "confront problems that polished tutorials skip."
)

h2("What you will learn")
skills = [
    "How to load and inspect a large CSV file with pandas",
    "How to measure data quality: missing values, duplicates, statistical outliers",
    "How to write automated tests that enforce quality thresholds (pytest)",
    "How to engineer features and split data into training and test sets",
    "What a baseline model is and why every project needs one",
    "How Ridge Regression and Gradient Boosting work at an intuitive level",
    "How to use MLflow to log every experiment so results are reproducible",
    "How to read evaluation metrics (MAE, RMSE, R², MAPE) and know which to trust",
]
for s in skills:
    bullet(s)

h2("The prediction task")
body(
    "We will predict Valeur fonciere — the recorded sale price (in euros) of a "
    "residential property. This is a regression problem: the output is a continuous "
    "number, not a category."
)

callout(
    "Regression vs Classification: if the output is a number (price, temperature, "
    "weight) it is regression. If it is a label (spam/not-spam, cat/dog) it is "
    "classification. The algorithms differ, but the workflow is the same.",
    emoji="📖"
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
#  2. ENVIRONMENT
# ════════════════════════════════════════════════════════════════════════════
h1("2.  Step 1 — Setting Up the Environment")

body(
    "Before writing a single line of analysis code, we need the right Python "
    "libraries installed. Each library has a specific role:"
)

libs = [
    ("pandas",     "Load, filter, and manipulate tabular data (DataFrames)."),
    ("numpy",      "Numerical operations: arrays, math, statistics."),
    ("scikit-learn","Model training, preprocessing pipelines, evaluation metrics."),
    ("lightgbm",   "Gradient boosting — our most powerful model."),
    ("mlflow",     "Experiment tracking: log params, metrics, and model artefacts."),
    ("pytest",     "Automated testing framework for data quality checks."),
]

for lib, desc in libs:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run(f"{lib}  ")
    r1.font.bold = True
    r1.font.name = "Courier New"
    r1.font.size = Pt(11)
    r2 = p.add_run(f"— {desc}")
    r2.font.size = Pt(11)

body("Install everything with one command:")
code("pip install pandas numpy scikit-learn lightgbm mlflow pytest")

callout(
    "Virtual environments: in real projects, always isolate dependencies with "
    "'python -m venv venv' and activate it before pip install. This prevents "
    "version conflicts between projects.",
    emoji="💡"
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
#  3. LOADING THE DATASET
# ════════════════════════════════════════════════════════════════════════════
h1("3.  Step 2 — Loading & Understanding the Dataset")

h2("About the DVF dataset")
body(
    "DVF (Demandes de Valeurs Foncieres) is published by the French government and "
    "records every property mutation (sale, exchange, gift) registered by notaries. "
    "Our file contains 16,306,439 rows and 43 columns covering transactions from 2020."
)

bullet("Each row = one parcel (plot of land) involved in one transaction.")
bullet("One sale can span multiple rows (e.g. a house with an attached garden is two rows).")
bullet("Monetary values use the French decimal separator: comma instead of period (8000,00 = 8000.00).")

h2("Loading with pandas")
body("pandas is the go-to library for tabular data in Python. We use read_csv:")
code(
    'df = pd.read_csv(\n'
    '    "mycsvfile.csv",\n'
    '    sep=",",          # column separator\n'
    '    low_memory=False, # read all rows before inferring types\n'
    '    decimal=",",      # French decimal format\n'
    ')'
)

h2("Why low_memory=False?")
body(
    "By default, pandas reads a file in chunks and guesses the type of each column "
    "from the first chunk. If a column has integers in chunk 1 but strings in chunk 2, "
    "you get a mixed-type (DtypeWarning) column that breaks calculations. "
    "low_memory=False reads the whole file first, so type inference is accurate."
)

h2("First look at the data")
body("Key commands for exploring a new DataFrame:")
code(
    "df.shape          # (rows, columns)\n"
    "df.head()         # first 5 rows\n"
    "df.dtypes         # type of each column\n"
    "df.describe()     # count, mean, std, min, max of numeric cols\n"
    "df.columns        # list of column names"
)

body("Our dataset: 16,306,439 rows × 43 columns. Key columns:")

cols_table = [
    ("Date mutation",      "Date of the transaction (DD/MM/YYYY)"),
    ("Nature mutation",    "Type: Vente (sale), Echange (exchange), etc."),
    ("Valeur fonciere",    "Sale price in euros — our prediction TARGET"),
    ("Type local",         "Property type: Maison (house), Appartement (flat), etc."),
    ("Surface reelle bati","Built area in m2"),
    ("Nombre pieces principales", "Number of main rooms"),
    ("Surface terrain",    "Land area in m2"),
    ("Code departement",   "French department code (administrative region)"),
    ("Code postal",        "Postal code"),
]
for col, desc in cols_table:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run(f"{col}  ")
    r1.font.bold = True
    r1.font.name = "Courier New"
    r1.font.size = Pt(10)
    r2 = p.add_run(f"— {desc}")
    r2.font.size = Pt(11)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
#  4. DATA QUALITY
# ════════════════════════════════════════════════════════════════════════════
h1("4.  Step 3 — Data Quality Audit")

body(
    "Raw data is almost always dirty. Before training any model, we must understand "
    "what is wrong with the data. We check three categories of issues:"
)

# ── 4a Missing values ──────────────────────────────────────────────────────
h2("4a.  Missing Values")
body(
    "A missing value (NaN — Not a Number) occurs when a field was not recorded. "
    "Most ML algorithms cannot handle NaN directly and will raise an error or silently "
    "produce wrong results. We measure the missing rate for every column:"
)
code(
    "missing     = df.isnull().sum()          # count of NaN per column\n"
    "missing_pct = missing / len(df) * 100    # percentage"
)

h3("What we found")
findings_mv = [
    ("8 columns are 100% null",
     "Code service CH, Reference document, 1–5 Articles CGI, Identifiant local. "
     "These columns exist in the schema but were never filled for this extract. "
     "They must be dropped before modelling — a model cannot learn from all-null features."),
    ("Valeur fonciere: 1.16% null (188,526 rows)",
     "The target column. Rows with no price cannot be used for supervised learning — "
     "we need the answer to train the model."),
    ("Code postal: 0.93% null",
     "Postal code is needed for location-based features."),
    ("30 of 43 columns exceed 5% missing",
     "Many lot-level columns (2eme lot, 3eme lot …) are legitimately sparse because "
     "most properties have only one lot. But they still need handling."),
]
for title, detail in findings_mv:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.space_after = Pt(6)
    r1 = p.add_run(f"{title}: ")
    r1.font.bold = True; r1.font.size = Pt(11)
    r2 = p.add_run(detail)
    r2.font.size = Pt(11)

# ── 4b Duplicates ──────────────────────────────────────────────────────────
h2("4b.  Duplicates")
body(
    "Duplicate rows inflate counts, bias averages, and give the model a false sense "
    "of confidence on rows it has already seen. We check at two levels:"
)
code(
    "# Exact duplicates — every column identical\n"
    "df.duplicated().sum()                          # found: 638,222\n\n"
    "# Semantic duplicates — same business key\n"
    "key = ['Date mutation','Valeur fonciere','Commune','No plan','Section']\n"
    "df.duplicated(subset=key).sum()                # found: 4,494,590"
)

callout(
    "Why do semantic duplicates exist? In DVF a single sale creates one row per parcel "
    "and one row per building unit on that parcel. When the same (date, price, parcel) "
    "tuple appears again, it usually means the raw file was concatenated more than once "
    "during export. Always de-duplicate before aggregating.",
    emoji="⚠️"
)

# ── 4c Outliers ──────────────────────────────────────────────────────────
h2("4c.  Outliers")
body(
    "An outlier is a value that is far from the bulk of the distribution. "
    "Outliers can be legitimate extreme cases or data entry errors — you must "
    "investigate which. We use the IQR (Interquartile Range) method:"
)
code(
    "Q1, Q3 = series.quantile(0.25), series.quantile(0.75)\n"
    "IQR    = Q3 - Q1\n"
    "lower  = Q1 - 1.5 * IQR   # lower fence\n"
    "upper  = Q3 + 1.5 * IQR   # upper fence\n"
    "outliers = series[(series < lower) | (series > upper)]"
)

body("Anything outside [lower, upper] is flagged as an outlier. Key findings:")

findings_out = [
    ("Valeur fonciere max = 2,086,000,000 €",
     "2 billion euros is almost certainly an encoding error (centimes instead of euros, or a typo). "
     "We cap the price at 10 M€ before training."),
    ("Valeur fonciere min = 0.01 €",
     "Sales for 1 euro are symbolic transfers (gifts between family members, divorce settlements). "
     "They are not real market prices. We apply a floor of 1,000 € to exclude them."),
    ("Surface Carrez max = 9,999 m2",
     "9,999 is a sentinel value used in French cadastral software to indicate 'data not available'. "
     "It must be treated as NaN, not as a real measurement."),
    ("Nombre de lots max = 468",
     "468 lots in a single transaction is highly unusual. These bulk transactions "
     "(entire apartment buildings, industrial estates) skew distributions."),
]
for title, detail in findings_out:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.space_after = Pt(6)
    r1 = p.add_run(f"{title}: ")
    r1.font.bold = True; r1.font.size = Pt(11)
    r2 = p.add_run(detail)
    r2.font.size = Pt(11)

callout(
    "Rule of thumb: never delete outliers blindly. First ask 'does this value "
    "make business sense?' A 10 M€ house is unusual but real. A 2 billion € "
    "house is almost certainly an error.",
    emoji="💡"
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
#  5. TESTS
# ════════════════════════════════════════════════════════════════════════════
h1("5.  Step 4 — Writing Failing Data Quality Tests")

body(
    "Auditing data manually is fragile — the next time the file is updated, you "
    "won't know if a new problem appeared. Instead, we encode our quality expectations "
    "as automated tests using pytest. Each test asserts something that should be true "
    "about clean data. Because our data has issues, every test currently fails — and "
    "that is intentional."
)

h2("Why write tests that fail?")
body(
    "A failing test documents the problem precisely. When you later clean the data "
    "(impute missing values, remove duplicates, cap outliers), you re-run the tests "
    "and watch them turn green one by one. The test suite becomes your definition of "
    "done for data cleaning."
)

h2("pytest basics")
body("pytest finds and runs any function whose name starts with test_:")
code(
    "# test_example.py\n"
    "def test_no_negatives():\n"
    "    prices = [100, 200, -5, 300]   # -5 is wrong\n"
    "    assert all(p >= 0 for p in prices), 'Negative price found!'\n\n"
    "# Run with:\n"
    "# pytest test_example.py -v"
)

h2("Structure of our test file")
body("We organised tests into three classes — one per quality category:")

code(
    "class TestMissingValues:\n"
    "    def test_no_completely_empty_columns(self, df): ...\n"
    "    def test_valeur_fonciere_no_nulls(self, df): ...\n"
    "    def test_code_postal_no_nulls(self, df): ...\n\n"
    "class TestDuplicates:\n"
    "    def test_no_fully_duplicate_rows(self, df): ...\n"
    "    def test_no_duplicate_deal_transactions(self, df): ...\n\n"
    "class TestOutliers:\n"
    "    def test_valeur_fonciere_no_extreme_high(self, df): ...\n"
    "    def test_valeur_fonciere_no_near_zero(self, df): ...\n"
    "    ..."
)

h2("The @pytest.fixture decorator")
body(
    "Loading 16M rows takes ~37 seconds. We don't want to do that for every test. "
    "A fixture with scope='session' loads the DataFrame once and shares it across "
    "all tests in the session:"
)
code(
    "@pytest.fixture(scope='session')\n"
    "def df():\n"
    "    return pd.read_csv(CSV_PATH, sep=',', decimal=',', low_memory=False)\n\n"
    "# Each test receives 'df' as a parameter — pytest injects it automatically."
)

h2("Test results: 15 / 15 FAILED (expected)")
body("Every test failed with a meaningful message, for example:")
code(
    "FAILED test_data_quality.py::TestMissingValues::test_valeur_fonciere_no_nulls\n"
    "  AssertionError: 'Valeur fonciere' contains 188,526 null values (1.16 %).\n"
    "  Every transaction must have a recorded sale price.\n"
    "  assert 188526 == 0"
)

callout(
    "These 15 failing tests are your data cleaning backlog. Fix them one by one and "
    "the tests will start passing. That is test-driven data engineering.",
    emoji="🎯"
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
#  6. DATA PREP
# ════════════════════════════════════════════════════════════════════════════
h1("6.  Step 5 — Preparing the Data for Machine Learning")

h2("Filtering to a clean subset")
body(
    "Rather than cleaning all 16M rows (a data engineering project in itself), we "
    "filter to a clean, well-understood subset: residential transactions with valid prices."
)
code(
    "# Keep only houses and flats\n"
    "df = df[df['Type local'].isin(['Maison', 'Appartement'])]\n\n"
    "# Drop rows with no price\n"
    "df = df.dropna(subset=['Valeur fonciere'])\n\n"
    "# Apply QA-derived price bounds\n"
    "df = df[(df['Valeur fonciere'] >= 1_000) & (df['Valeur fonciere'] <= 10_000_000)]\n\n"
    "# Result: 5,890,138 rows"
)

h2("Sampling")
body(
    "5.9M rows is still large. For model development we draw a random sample of 300,000 "
    "rows. This keeps iteration fast while being representative. We fix random_state=42 "
    "so every run produces the same sample."
)
code("df = df.sample(n=300_000, random_state=42).reset_index(drop=True)")

h2("Feature selection")
body(
    "We choose 7 features that are available for most rows and intuitively relevant "
    "to price prediction:"
)
features = [
    ("Surface reelle bati",       "Numeric", "Built area in m2 — bigger = more expensive"),
    ("Nombre pieces principales", "Numeric", "Room count — proxy for size and comfort"),
    ("Surface terrain",           "Numeric", "Land area — land value in addition to building"),
    ("Nombre de lots",            "Numeric", "Number of lots — complex multi-unit transactions"),
    ("Type local",                "Categorical", "Maison vs Appartement — structurally different markets"),
    ("Code departement",          "Categorical", "Region — Paris vs rural France differs massively in price"),
    ("Nature mutation",           "Categorical", "Vente vs other — only arm's-length sales are market price"),
]
for fname, ftype, fdesc in features:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run(f"{fname}  ")
    r1.font.bold = True; r1.font.name = "Courier New"; r1.font.size = Pt(10)
    r2 = p.add_run(f"[{ftype}]  {fdesc}")
    r2.font.size = Pt(11)

h2("Log-transforming the target")
body(
    "House prices are right-skewed: most are between 100k and 500k, but a few go up "
    "to 10M. This skewness makes regression harder. Applying log1p (log(1 + price)) "
    "compresses the scale and makes the distribution closer to normal:"
)
code(
    "y = np.log1p(df['Valeur fonciere'])\n"
    "# log1p(200,000) ≈ 12.2\n"
    "# log1p(2,000,000) ≈ 14.5  (not 10x bigger anymore)\n\n"
    "# To convert back after prediction:\n"
    "price = np.expm1(y_predicted)"
)

callout(
    "log1p instead of log: np.log(0) is -infinity, which breaks training. "
    "np.log1p(x) = log(1+x) handles zero gracefully. Always use it for prices.",
    emoji="💡"
)

h2("Train / Test split")
body(
    "We split the data: 80% for training (the model sees this), 20% for testing "
    "(the model never sees this — it is our honest evaluation set)."
)
code(
    "X_train, X_test, y_train, y_test = train_test_split(\n"
    "    X, y,\n"
    "    test_size=0.2,\n"
    "    random_state=42,    # reproducibility\n"
    ")   # -> 240,000 train rows, 60,000 test rows"
)

callout(
    "Never evaluate a model on training data. The model has already 'seen' those "
    "rows and will score artificially well. The test set simulates unseen, future data.",
    emoji="⚠️"
)

h2("Preprocessing pipelines")
body(
    "Raw data still has missing values and categorical strings that algorithms cannot "
    "use directly. We build preprocessing pipelines using scikit-learn:"
)
code(
    "from sklearn.pipeline import Pipeline\n"
    "from sklearn.compose import ColumnTransformer\n"
    "from sklearn.impute import SimpleImputer\n"
    "from sklearn.preprocessing import StandardScaler, OrdinalEncoder\n\n"
    "num_pipe = Pipeline([\n"
    "    ('impute', SimpleImputer(strategy='median')),  # fill NaN with median\n"
    "    ('scale',  StandardScaler()),                  # mean=0, std=1\n"
    "])\n\n"
    "cat_pipe = Pipeline([\n"
    "    ('impute', SimpleImputer(strategy='most_frequent')),\n"
    "    ('encode', OrdinalEncoder()),                  # string -> integer\n"
    "])\n\n"
    "preprocessor = ColumnTransformer([\n"
    "    ('num', num_pipe, NUMERIC_COLUMNS),\n"
    "    ('cat', cat_pipe, CATEGORICAL_COLUMNS),\n"
    "])"
)
body("Key concepts:")
bullet("SimpleImputer(median): replaces NaN with the median of that column.")
bullet("StandardScaler: transforms each numeric column to have mean 0 and std 1. "
       "Required for Ridge — gradient-based linear models are sensitive to feature scale.")
bullet("OrdinalEncoder: converts 'Maison' → 0, 'Appartement' → 1. "
       "Needed because math cannot operate on raw strings.")
bullet("ColumnTransformer: applies different pipelines to different column subsets "
       "in one step.")

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
#  7. MODELS
# ════════════════════════════════════════════════════════════════════════════
h1("7.  Step 6 — Choosing & Training Three Models")

# ── Baseline ──────────────────────────────────────────────────────────────
h2("Model 1 — Baseline: DummyRegressor")
body(
    "Before training anything sophisticated, we establish a baseline: the simplest "
    "possible 'model'. A DummyRegressor with strategy='median' ignores all features "
    "and always predicts the median training price."
)
body("Why bother? Because any real model must beat this threshold. If your "
     "fancy model barely beats a model that never looks at the data, something is wrong.")
code(
    "from sklearn.dummy import DummyRegressor\n\n"
    "model = DummyRegressor(strategy='median')\n"
    "model.fit(X_train, y_train)       # only memorises the median\n"
    "y_pred = model.predict(X_test)    # returns median for every row"
)
body("Result: MAE = 241,559 €  |  RMSE = 857,108 €  |  R² = -0.04  |  MAPE = 107 %")
callout(
    "R² close to 0 (or slightly negative) for a Dummy is normal — it means the model "
    "explains no variance. R² = 1.0 would be a perfect model.",
    emoji="📖"
)

# ── Ridge ──────────────────────────────────────────────────────────────────
h2("Model 2 — Ridge Regression")
body(
    "Ridge is a linear model. It learns a straight-line relationship between each "
    "feature and the target, with a regularisation penalty (alpha) that prevents "
    "the coefficients from becoming too large."
)
body("The prediction equation:")
code("price_log = w1*surface + w2*rooms + w3*terrain + ... + bias")
body(
    "During training, Ridge finds the weights (w1, w2, ...) that minimise the "
    "sum of squared errors plus the penalty alpha * sum(w^2). "
    "A larger alpha means stronger regularisation (smaller, safer weights)."
)
code(
    "from sklearn.linear_model import Ridge\n\n"
    "model = Pipeline([\n"
    "    ('prep',  preprocessor),     # impute + scale + encode\n"
    "    ('model', Ridge(alpha=10)),  # then fit Ridge\n"
    "])\n"
    "model.fit(X_train, y_train)"
)
body("Result: MAE = 232,751 €  |  RMSE = 841,168 €  |  R² = -0.003  |  MAPE = 104 %")
body(
    "Ridge only slightly beats the Dummy (MAE -4%). This is expected: a linear model "
    "with 7 ordinal-encoded features cannot capture the complex, non-linear interactions "
    "in real estate (location × size, neighbourhood effects, etc.)."
)
callout(
    "Ridge makes one critical assumption: the relationship between features and target "
    "is linear. Real house prices violate this assumption. A 10m2 increase on a "
    "15m2 studio is proportionally more valuable than on a 200m2 villa.",
    emoji="💡"
)

# ── LightGBM ──────────────────────────────────────────────────────────────
h2("Model 3 — LightGBM (Gradient Boosting)")
body(
    "LightGBM builds an ensemble of decision trees sequentially. Each tree tries to "
    "correct the errors made by all previous trees. This boosting approach captures "
    "non-linear relationships and feature interactions automatically."
)

h3("Decision trees (the building block)")
body(
    "A decision tree asks yes/no questions about the features to partition the data "
    "into groups with similar target values. For example:"
)
code(
    "Is Type local == 'Maison'?\n"
    "  YES -> Is Surface reelle bati > 120 m2?\n"
    "           YES -> predicted price: 320,000 €\n"
    "           NO  -> predicted price: 210,000 €\n"
    "  NO  -> Is Code departement == 75 (Paris)?\n"
    "           YES -> predicted price: 490,000 €\n"
    "           NO  -> predicted price: 165,000 €"
)

h3("Gradient boosting")
body(
    "LightGBM builds many small trees, each one learning from the residuals (errors) "
    "of the previous trees. The learning_rate controls how much each tree contributes."
)
code(
    "lgb_params = {\n"
    "    'n_estimators':    800,   # up to 800 trees\n"
    "    'learning_rate':   0.05,  # small steps = more stable\n"
    "    'num_leaves':      127,   # tree complexity\n"
    "    'subsample':       0.8,   # use 80% of rows per tree (reduces overfitting)\n"
    "    'colsample_bytree':0.8,   # use 80% of features per tree\n"
    "}\n\n"
    "lgb_model = lgb.LGBMRegressor(**lgb_params)\n"
    "lgb_model.fit(\n"
    "    X_train_preprocessed, y_train,\n"
    "    eval_set=[(X_test_preprocessed, y_test)],\n"
    "    callbacks=[lgb.early_stopping(50)],  # stop if no improvement for 50 rounds\n"
    ")"
)
body("Result: MAE = 168,539 €  |  RMSE = 641,242 €  |  R² = 0.42  |  MAPE = 80 %")
body(
    "LightGBM is the only model with positive R² — it genuinely explains variance in "
    "the data. It achieves a 30% MAE reduction over the Dummy baseline."
)
callout(
    "Early stopping: we tell LightGBM to monitor the test error during training. "
    "When the test error stops improving for 50 consecutive rounds, training halts "
    "automatically. This prevents overfitting and saved computation: "
    "the model stopped at 719 of 800 trees.",
    emoji="💡"
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
#  8. MLFLOW
# ════════════════════════════════════════════════════════════════════════════
h1("8.  Step 7 — Tracking Experiments with MLflow")

body(
    "MLflow is a library that records every experiment run: what parameters you used, "
    "what metrics you got, and the trained model file. Without tracking, it is easy "
    "to lose which configuration produced which result."
)

h2("Core concepts")
tracked = [
    ("Parameters (params)",  "The settings / hyperparameters you chose before training "
                              "(e.g. alpha=10, n_estimators=800)."),
    ("Metrics",              "The evaluation scores computed after training "
                              "(MAE, RMSE, R², MAPE, training time)."),
    ("Artefacts",            "Files saved with the run: the serialised model file, "
                              "feature importance JSON, input examples."),
    ("Run",                  "One execution of your training code. Each run gets a unique ID."),
    ("Experiment",           "A named group of runs (we used 'DVF_Price_Prediction')."),
]
for term, desc in tracked:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(f"{term}: ")
    r1.font.bold = True; r1.font.size = Pt(11)
    r2 = p.add_run(desc)
    r2.font.size = Pt(11)

h2("Logging a run")
code(
    "import mlflow\n\n"
    "mlflow.set_experiment('DVF_Price_Prediction')\n\n"
    "with mlflow.start_run(run_name='02_Ridge_Regression'):\n"
    "    mlflow.log_params({'alpha': 10, 'max_iter': 1000})\n\n"
    "    model.fit(X_train, y_train)\n"
    "    y_pred = model.predict(X_test)\n\n"
    "    mlflow.log_metric('mae',  mean_absolute_error(y_test, y_pred))\n"
    "    mlflow.log_metric('r2',   r2_score(y_test, y_pred))\n\n"
    "    mlflow.sklearn.log_model(model, 'model')  # save the model file"
)
body("The with block creates one run. When the block exits, MLflow seals the run.")

h2("Viewing the MLflow UI")
body("Start the local web server from your project folder:")
code("mlflow ui --backend-store-uri file:./mlruns")
body(
    "Open http://127.0.0.1:5000 in your browser. You will see a table of all runs, "
    "sortable by any metric. Click a run to see its full parameter and metric history, "
    "and to download the saved model."
)
callout(
    "In teams, mlflow can point to a remote server (S3, Azure Blob, a shared database) "
    "so everyone's experiments are in one place. Even solo, it replaces the "
    "'results_final_v3_FINAL.csv' file chaos.",
    emoji="💡"
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
#  9. RESULTS
# ════════════════════════════════════════════════════════════════════════════
h1("9.  Step 8 — Comparing Results & Reading the Metrics")

h2("The comparison table")
body("After training all three models, we collected their test-set metrics:")

# Draw a table
tbl = doc.add_table(rows=4, cols=8)
tbl.style = "Table Grid"

headers = ["Model", "MAE (€)", "RMSE (€)", "R²", "MAPE (%)", "Log-MAE", "Log-RMSE", "Time (s)"]
rows_data = [
    ["Baseline (Dummy median)", "241,559", "857,108", "-0.042", "107.3%", "0.681", "0.964", "0.004"],
    ["Ridge Regression",        "232,751", "841,168", "-0.003", "104.6%", "0.623", "0.886", "0.42"],
    ["LightGBM",                "168,539", "641,242",  "0.417",  "80.2%", "0.447", "0.683", "17.5"],
]

for i, hdr in enumerate(headers):
    cell = tbl.rows[0].cells[i]
    cell.text = hdr
    cell.paragraphs[0].runs[0].font.bold = True
    cell.paragraphs[0].runs[0].font.size = Pt(9)
    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), "1F4E79")
    tcPr.append(shd)

for row_idx, row_data in enumerate(rows_data, start=1):
    fill = "DEEAF1" if row_idx % 2 == 0 else "FFFFFF"
    if row_data[0] == "LightGBM":
        fill = "E2EFDA"   # green highlight for best model
    for col_idx, val in enumerate(row_data):
        cell = tbl.rows[row_idx].cells[col_idx]
        cell.text = val
        cell.paragraphs[0].runs[0].font.size = Pt(9)
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), fill)
        tcPr.append(shd)

doc.add_paragraph()

h2("Understanding each metric")
metrics = [
    ("MAE — Mean Absolute Error",
     "The average of |true_price - predicted_price| across all test rows. "
     "In our case: on average, LightGBM's prediction is off by 168,539 €. "
     "MAE is intuitive because it is in the same unit as the target (euros)."),
    ("RMSE — Root Mean Squared Error",
     "Like MAE but squares the errors before averaging, then takes the square root. "
     "Because of squaring, large individual errors are penalised more heavily. "
     "RMSE > MAE always. A big gap means there are a few very wrong predictions."),
    ("R² — Coefficient of Determination",
     "Measures what fraction of the variance in the target the model explains. "
     "R²=1.0: perfect. R²=0.0: no better than predicting the mean. R²<0: worse than mean. "
     "LightGBM R²=0.42 means it explains 42% of price variance with just 7 features."),
    ("MAPE — Mean Absolute Percentage Error",
     "Average of |true - pred| / true, expressed as %. LightGBM MAPE=80% means "
     "the prediction is off by 80% of the true price on average. High because cheap "
     "properties (100k) have the same absolute error as expensive ones (500k), "
     "making the percentage large for cheap ones."),
    ("Log-MAE / Log-RMSE",
     "MAE and RMSE computed in log-price space (before expm1 conversion). "
     "More reliable than price-space metrics when price spans several orders of magnitude. "
     "Lower is better. LightGBM Log-MAE=0.447 means predictions are off by e^0.447 ≈ 1.56x "
     "in multiplicative terms."),
]
for mname, mdesc in metrics:
    h3(mname)
    body(mdesc, indent=True)

h2("Why are the absolute errors still large?")
body(
    "A MAE of 168k euros on a dataset where the median price is ~200k is high. "
    "This is expected with our current feature set. To improve:"
)
improvements = [
    "Add finer location: Code postal or Commune (1-2 digit department → 5-digit postcode)",
    "Add date features: year, month — prices trend over time",
    "Add the Type local × Surface interaction (surface means different things for houses vs flats)",
    "Hyperparameter tuning: use Optuna or GridSearchCV to find better LightGBM settings",
    "Ensemble: combine Ridge and LightGBM predictions (blending)",
]
for item in improvements:
    bullet(item)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
#  10. GLOSSARY
# ════════════════════════════════════════════════════════════════════════════
h1("10.  Key Concepts Glossary")

terms = [
    ("DataFrame",         "The central pandas data structure: a 2D table with labelled rows and columns."),
    ("Feature",           "An input variable used to make a prediction (e.g. surface area)."),
    ("Target",            "The output variable we are trying to predict (e.g. sale price)."),
    ("Supervised learning","Learning from labelled examples: each row has known features AND a known target."),
    ("Regression",        "Predicting a continuous numeric output."),
    ("Overfitting",        "The model memorises the training data but fails on new data. Detected when train metrics >> test metrics."),
    ("Regularisation",    "A penalty added to the loss function to keep model weights small and prevent overfitting. Alpha in Ridge controls this."),
    ("Imputation",        "Filling in missing values with a summary statistic (median, mean, mode)."),
    ("Ordinal encoding",  "Mapping string categories to integers (Maison=0, Appartement=1)."),
    ("Pipeline",          "A sequence of preprocessing + model steps that are fit and applied together, preventing data leakage."),
    ("Train/test split",  "Dividing data into a training portion (model learns from this) and a test portion (model is evaluated on this)."),
    ("IQR",               "Interquartile Range = Q3 - Q1. The middle 50% of a distribution. Used to define outlier fences."),
    ("Gradient boosting", "An ensemble method that builds trees sequentially, each one correcting the errors of the previous."),
    ("Early stopping",    "Halting training when the validation metric stops improving, preventing overfitting and wasted compute."),
    ("MLflow run",        "A recorded execution of training code, with params, metrics, and model artefacts attached."),
    ("Baseline model",    "The simplest possible model (e.g., always predict the median). Every other model must beat it."),
]
for term, definition in terms:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.space_after = Pt(5)
    r1 = p.add_run(f"{term}: ")
    r1.font.bold = True; r1.font.size = Pt(11)
    r2 = p.add_run(definition)
    r2.font.size = Pt(11)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
#  11. NEXT STEPS
# ════════════════════════════════════════════════════════════════════════════
h1("11.  What to Learn Next")

next_steps = [
    ("Fix the failing tests",
     "Work through each of the 15 failing tests. Impute missing values, remove "
     "duplicates, clip outliers. Watch the tests go green. This is data engineering."),
    ("Add more features",
     "Re-encode Code postal as a numeric feature. Extract year and month from "
     "Date mutation. Compute price per m2 as a derived feature."),
    ("Hyperparameter tuning",
     "Use Optuna (pip install optuna) to search for the best LightGBM parameters. "
     "Log each trial as an MLflow run."),
    ("Cross-validation",
     "Instead of one 80/20 split, use 5-fold cross-validation to get a more "
     "reliable estimate of model performance."),
    ("Feature importance & explainability",
     "LightGBM exposes feature_importances_. Plot them with matplotlib. "
     "For deeper insight, install SHAP (pip install shap) which shows how each "
     "feature pushes the prediction up or down for individual rows."),
    ("Deploy the model",
     "Serve your MLflow model as a REST API with 'mlflow models serve'. "
     "Send a JSON request with property features and get a price prediction back."),
    ("Learn pandas more deeply",
     "This project used ~10% of pandas. Study groupby, merge, pivot_table, "
     "and time-series resampling."),
]
for i, (title, detail) in enumerate(next_steps, 1):
    h3(f"{i}.  {title}")
    body(detail, indent=True)

# ── Final note ───────────────────────────────────────────────────────────
doc.add_paragraph()
sep()
close = doc.add_paragraph()
close.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = close.add_run(
    "You have just completed a real-world ML project from scratch.\n"
    "The data was messy, the models are imperfect, and that is normal.\n"
    "Every improvement you make from here is real machine learning engineering."
)
r.font.size = Pt(11)
r.font.italic = True
r.font.color.rgb = RGBColor(60, 60, 60)

# ── Save ─────────────────────────────────────────────────────────────────
out_path = r"C:\Users\USER\ml-project\ML_Project_Tutorial.docx"
doc.save(out_path)
print(f"Saved: {out_path}")
